"""
AdSphere RTB System — Concise Project Documentation Generator
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ARTIFACT_DIR = "/Users/sakshamm/.gemini/antigravity/brain/69924b2f-3d40-4059-8504-fdceab51feb7"
ARCH_IMG  = "docs/architecture_diagram.png"
DB_IMG    = os.path.join(ARTIFACT_DIR, "database_schema_1781787748656.png")
TECH_IMG  = os.path.join(ARTIFACT_DIR, "tech_stack_diagram_1781787760285.png")
FLOW_IMG  = os.path.join(ARTIFACT_DIR, "rtb_auction_flow_1781787792041.png")
DASH_IMG  = os.path.join(ARTIFACT_DIR, "dashboard_screenshot_1781787802968.png")
LOG_IMG   = os.path.join(ARTIFACT_DIR, "bid_logs_screenshot_1781787853532.png")

NAVY   = RGBColor(0x1E, 0x3A, 0x5F)
PURPLE = RGBColor(0x5B, 0x21, 0xB6)
BODY   = RGBColor(0x37, 0x4A, 0x5E)
DARK   = RGBColor(0x1E, 0x29, 0x3B)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    # underline rule
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), "5B21B6")
    pb.append(bot); pPr.append(pb)

def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = PURPLE
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)

def para(doc, text):
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.size = Pt(10); r.font.color.rgb = BODY
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(14)

def bul(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(10); r.font.color.rgb = BODY
    p.paragraph_format.space_after = Pt(2)

def img(doc, path, width=5.8, caption=""):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.runs[0]
        r.italic = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        cp.paragraph_format.space_after = Pt(6)

def table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        c.text = h
        set_cell_bg(c, "1E3A5F")
        r = c.paragraphs[0].runs[0]
        r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # data rows
    for idx, row in enumerate(rows):
        cells = tbl.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
            r = cells[i].paragraphs[0].runs[0]
            r.font.size = Pt(9); r.font.color.rgb = DARK
            if idx % 2 == 0:
                set_cell_bg(cells[i], "F1F5F9")
    if col_widths:
        for ri in range(len(tbl.rows)):
            for ci, w in enumerate(col_widths):
                tbl.rows[ri].cells[ci].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ════════════════════════════════════════════════════════════════════════════
doc = Document()

sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Cm(1.8)
sec.left_margin = sec.right_margin = Cm(2.2)

sty = doc.styles["Normal"]
sty.font.name = "Calibri"
sty.font.size = Pt(10)
sty.font.color.rgb = DARK

# ── COVER PAGE ───────────────────────────────────────────────────────────────
for _ in range(3): doc.add_paragraph()

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("AdSphere"); r.bold = True; r.font.size = Pt(32); r.font.color.rgb = NAVY

s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Real-Time Bidding (RTB) Advertising System")
r.bold = True; r.font.size = Pt(15); r.font.color.rgb = PURPLE

doc.add_paragraph()
for label, val in [("Type", "College System Design Capstone Project"),
                   ("Stack", "FastAPI · Python · PostgreSQL/SQLite · Redis"),
                   ("Date",  "June 2026")]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{label}:  "); r1.bold = True; r1.font.size = Pt(10)
    r2 = p.add_run(val); r2.font.size = Pt(10); r2.font.color.rgb = BODY

doc.add_page_break()

# ── 1. PROBLEM STATEMENT ─────────────────────────────────────────────────────
h1(doc, "1.  Problem Statement")
para(doc,
    "In today's digital world, billions of ads are shown every second across websites and apps. "
    "Traditional advertising showed the same ad to every user regardless of age, location, or interest — "
    "wasting advertiser budgets and frustrating users with irrelevant content. There was no real-time "
    "control, no smart targeting, and no protection against bot traffic and fake clicks.")

h2(doc, "Core Problems")
for b in ["Ads were irrelevant — wrong ads shown to the wrong people",
          "Ad budgets were wasted on users with zero interest",
          "No real-time control — advertisers had no way to set smart rules",
          "Fraud was widespread — fake clicks and bot traffic drained money"]:
    bul(doc, b)

# ── 2. PROPOSED SOLUTION ─────────────────────────────────────────────────────
h1(doc, "2.  Proposed Solution")
para(doc,
    "AdSphere is a full-stack Real-Time Bidding (RTB) platform analogous to Google Ads or Meta Audience Network, "
    "built from scratch. When a user visits a page, a silent auction runs in under 15ms — multiple advertisers "
    "compete, the smartest bid wins, and the right ad is shown to the right person instantly.")

table(doc,
    headers=["Feature", "Description"],
    rows=[
        ["Campaign Management",   "Create campaigns with budgets, bids, and targeting rules"],
        ["Smart Targeting",       "Filter by age, location, and interest overlap"],
        ["Real-Time Auction",     "Sub-15ms auction engine with full scoring"],
        ["Budget Enforcement",    "Atomic cache-based deductions synced to DB"],
        ["Fraud Protection",      "IP blacklist + sliding-window rate limiting"],
        ["Click & View Tracking", "1×1 GIF pixel + HTTP 307 click redirect"],
        ["Analytics Dashboard",   "Live revenue, impressions, clicks, and CTR charts"],
    ],
    col_widths=[5.0, 12.0],
)

# ── 3. SYSTEM ARCHITECTURE ───────────────────────────────────────────────────
h1(doc, "3.  System Architecture")
para(doc,
    "The diagram below shows how all components connect — from the client browser through fraud checks, "
    "the bidding engine, budget management, and back to the user as a winning ad with tracking URLs.")

img(doc, ARCH_IMG, width=5.6, caption="Figure 1 — AdSphere RTB System Architecture")

# ── 4. MODULE DESCRIPTION ────────────────────────────────────────────────────
h1(doc, "4.  Module Description")
table(doc,
    headers=["Module", "File", "Responsibility"],
    rows=[
        ["FastAPI API Gateway",  "src/api/",                          "HTTP entry point — routes all REST requests"],
        ["Fraud Guard",          "src/fraud/fraud_detector.py",        "IP blacklist check + rate-limit enforcement"],
        ["Bidding Engine",       "src/bidding/bidding_engine.py",      "Orchestrates full auction lifecycle"],
        ["Ranking Engine",       "src/bidding/ranking_engine.py",      "Jaccard similarity + Laplace CTR scoring"],
        ["Budget Manager",       "src/bidding/budget_manager.py",      "Atomic spend deduction (cache → DB)"],
        ["Cache Manager",        "src/cache/redis_cache.py",           "Redis client with in-memory fallback"],
        ["Analytics Processor",  "src/analytics/analytics_processor.py","Background impression/click event handler"],
    ],
    col_widths=[4.0, 5.5, 7.5],
)

# ── 5. DATABASE DESIGN ───────────────────────────────────────────────────────
h1(doc, "5.  Database Design")
para(doc,
    "AdSphere uses 5 SQLAlchemy ORM tables. Development uses SQLite (zero setup); "
    "production connects to PostgreSQL.")

img(doc, DB_IMG, width=5.4, caption="Figure 2 — AdSphere Database Schema (ER Diagram)")

table(doc,
    headers=["Table", "Purpose", "Key Columns"],
    rows=[
        ["campaigns",   "Ad campaign definitions",        "id, name, budget, current_spend, bid_amount, target_*, is_active"],
        ["users",       "Audience demographic profiles",  "id, name, age, location, interests"],
        ["bid_history", "Auction audit trail",            "campaign_id (FK), user_id (FK), score, status, reason"],
        ["analytics",   "Daily performance metrics",      "campaign_id (FK), impressions, clicks, revenue"],
        ["fraud_logs",  "Security event records",         "ip_address, reason, score, request_data"],
    ],
    col_widths=[3.5, 4.5, 9.0],
)

# ── 6. TECHNOLOGY STACK ──────────────────────────────────────────────────────
h1(doc, "6.  Technology Stack")
img(doc, TECH_IMG, width=4.8, caption="Figure 3 — AdSphere Technology Stack")

table(doc,
    headers=["Layer", "Technology", "Purpose"],
    rows=[
        ["Backend API",    "FastAPI (Python ≥ 3.9)",    "Async REST API with auto-generated Swagger docs"],
        ["ASGI Server",    "Uvicorn",                   "Lightweight async production server"],
        ["ORM",            "SQLAlchemy 2.0",            "Python-to-SQL table mapping, no raw queries"],
        ["DB (Dev)",       "SQLite",                    "Zero-setup file-based database for demos"],
        ["DB (Prod)",      "PostgreSQL",                "Production-grade relational database"],
        ["Cache",          "Redis / In-Memory fallback","Sub-ms budget checks and rate limiting"],
        ["Validation",     "Pydantic v2",               "Type-safe request/response payload models"],
        ["Frontend",       "HTML · CSS · JS · Chart.js","Dashboard UI with live charts"],
        ["Event Streaming","Mock Kafka Pipeline",       "Simulated producer/consumer event stream"],
    ],
    col_widths=[3.5, 4.5, 9.0],
)

# ── 7. IMPLEMENTATION DETAILS ────────────────────────────────────────────────
h1(doc, "7.  Implementation Details")

h2(doc, "Auction Scoring Formula")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Score  =  Bid Amount  ×  CTR  ×  Relevance Score")
r.bold = True; r.font.size = Pt(12); r.font.color.rgb = NAVY
p.paragraph_format.space_after = Pt(4)

h2(doc, "Relevance Score — Jaccard Similarity")
para(doc,
    "Interest overlap between user and campaign is measured using Jaccard Index = |Intersection| ÷ |Union|, "
    "then mapped to [0.1, 1.0] via:  Relevance = 0.1 + (Jaccard × 0.9).  "
    "Example: User {Tech, Gaming, Gadgets} ∩ Campaign {Tech, Gadgets} = 2/3 = 0.667 → Relevance = 0.70")

h2(doc, "Four-Stage Targeting Filter")
table(doc,
    headers=["Stage", "Filter", "Fail Reason Logged"],
    rows=[
        ["1", "Budget:  current_spend + bid_amount ≤ budget",           "FILTERED_BUDGET"],
        ["2", "Age:  target_age_min ≤ user.age ≤ target_age_max",       "FILTERED_TARGETING"],
        ["3", "Location:  'All' OR matches user city",                   "FILTERED_TARGETING"],
        ["4", "Interest:  ≥ 1 common interest between user & campaign",  "FILTERED_TARGETING"],
    ],
    col_widths=[1.5, 10.5, 5.0],
)

h2(doc, "Fraud Detection")
table(doc,
    headers=["Mechanism", "Threshold", "Action"],
    rows=[
        ["IP Blacklist",  "Exact match in blacklist",       "Block immediately, score = 1.0"],
        ["Rate Limiting", "> 10 requests / 10 seconds",     "Block request, score = 0.9"],
    ],
    col_widths=[4.0, 6.0, 7.0],
)

h2(doc, "Caching Strategy")
table(doc,
    headers=["Cache Key", "Data Stored", "TTL"],
    rows=[
        ["user:profile:{id}",       "Full user profile (JSON)",    "10 min"],
        ["campaign:spend:{id}",     "Current spend (float)",       "1 hour"],
        ["campaign:stats:{id}",     "Impressions + clicks",        "2 min"],
        ["rate_limit:{ip}",         "Request counter (int)",       "10 sec"],
    ],
    col_widths=[5.5, 6.5, 5.0],
)

h2(doc, "Key API Endpoints")
table(doc,
    headers=["Method", "Endpoint", "Description"],
    rows=[
        ["POST", "/auction/request",                              "⚡ Run a real-time bid auction"],
        ["POST", "/campaign",                                     "Create a new ad campaign"],
        ["GET",  "/campaigns",                                    "List all campaigns"],
        ["POST", "/user",                                         "Create a user profile"],
        ["GET",  "/analytics/summary",                            "Aggregated analytics data"],
        ["GET",  "/analytics/impression/{campaign_id}/{user_id}", "Impression tracking pixel (1×1 GIF)"],
        ["GET",  "/analytics/click/{campaign_id}/{user_id}",      "Click redirect (HTTP 307)"],
        ["GET",  "/history",                                      "Bid history audit log"],
        ["GET",  "/docs",                                         "Interactive Swagger UI"],
    ],
    col_widths=[2.0, 7.5, 7.5],
)

doc.add_page_break()

# ── 8. SCREENSHOTS ───────────────────────────────────────────────────────────
h1(doc, "8.  Screenshots")

h2(doc, "Overview Dashboard")
img(doc, DASH_IMG, width=5.6,
    caption="Figure 4 — KPI Cards (Revenue, Impressions, Clicks, CTR) + Campaign Spend Chart")

h2(doc, "Bid History & Fraud Log")
img(doc, LOG_IMG, width=5.6,
    caption="Figure 5 — Auction Bid History (WON/LOST/FILTERED) and Fraud Security Events")

h2(doc, "Auction Flow Diagram")
img(doc, FLOW_IMG, width=3.6,
    caption="Figure 6 — RTB Auction Flowchart")

doc.add_page_break()

# ── 9. FUTURE SCOPE ──────────────────────────────────────────────────────────
h1(doc, "9.  Future Scope")
table(doc,
    headers=["Enhancement", "Description"],
    rows=[
        ["ML Bid Optimization",    "Replace static formula with XGBoost/neural net for optimal bid prediction"],
        ["OpenRTB 2.6 Compliance", "Connect to real SSPs/DSPs using the industry-standard bid protocol"],
        ["Second-Price Auction",   "Implement Vickrey auction — winner pays 2nd-highest bid + $0.01"],
        ["Real Kafka Integration", "Replace mock pipeline with Apache Kafka for production event streaming"],
        ["Docker Deployment",      "Containerize all services (API, DB, Redis, Kafka) with docker-compose"],
        ["Multi-Format Ads",       "Add banner, video pre-roll, and native ad formats (CPM/CPC/CPV pricing)"],
        ["A/B Testing",            "Split ad creative traffic automatically and measure performance"],
        ["JWT Authentication",     "Multi-tenant login so each advertiser manages only their own campaigns"],
        ["IP Geolocation",         "Auto-detect user city from IP using MaxMind GeoIP — no manual input"],
        ["Real-Time Alerts",       "Email/Slack alerts when fraud spikes or budget exhaustion is near"],
    ],
    col_widths=[5.0, 12.0],
)

doc.add_paragraph()
closing = doc.add_paragraph("— End of Document —")
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = closing.runs[0]; r.bold = True; r.font.size = Pt(11); r.font.color.rgb = PURPLE

note = doc.add_paragraph("AdSphere RTB System  |  College Capstone Project  |  June 2026")
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = note.runs[0]; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

# ── SAVE ─────────────────────────────────────────────────────────────────────
out = "AdSphere_Project_Documentation.docx"
doc.save(out)
print(f"\n✅  Saved → {out}\n")
